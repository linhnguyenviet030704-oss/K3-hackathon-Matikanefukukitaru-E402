"""
rag/loaders.py — Document loaders for PDF, DOCX (.docx), DOC (.doc), and TXT files.

Supports recursive scanning of data/raw/ sub-folders.
  - .pdf   : via pypdf
  - .docx  : via python-docx (Office Open XML)
  - .doc   : via macOS textutil (binary Word 97-2003 format)
  - .txt   : plain read
"""

from __future__ import annotations

import logging
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class RawDocument:
    """A raw document loaded from disk before chunking."""

    text: str
    source_file: str          # Relative path from data/
    source_type: str          # "textbook" | "guideline" | "article"
    file_format: str          # "pdf" | "docx" | "txt"
    total_pages: int = 0
    extra: dict = field(default_factory=dict)


# ── PDF ───────────────────────────────────────────────────────────────────────

def load_pdf(path: Path, source_type: str = "article") -> RawDocument:
    """Load a PDF file and extract text page by page."""
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError("pypdf is required: pip install pypdf") from e

    reader = PdfReader(str(path))
    pages_text: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            pages_text.append(extracted)

    full_text = "\n\n".join(pages_text)
    logger.info("Loaded PDF: %s (%d pages)", path.name, len(reader.pages))

    return RawDocument(
        text=full_text,
        source_file=path.name,
        source_type=source_type,
        file_format="pdf",
        total_pages=len(reader.pages),
    )


# ── DOCX (.docx — Office Open XML) ───────────────────────────────────────────

def load_docx(path: Path, source_type: str = "textbook") -> RawDocument:
    """Load a .docx file using python-docx (requires Office Open XML format)."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("python-docx is required: pip install python-docx") from e

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    logger.info("Loaded DOCX: %s (%d paragraphs)", path.name, len(paragraphs))

    return RawDocument(
        text=full_text,
        source_file=path.name,
        source_type=source_type,
        file_format="docx",
    )


# ── DOC (.doc — binary Word 97-2003) ─────────────────────────────────────────

def load_doc(path: Path, source_type: str = "textbook") -> RawDocument:
    """
    Load a legacy binary .doc file.

    Strategy (tried in order):
      1. macOS textutil  — no extra install needed on macOS.
      2. python-docx     — works if the .doc is actually a disguised .docx.
      3. Raise a clear error with conversion instructions.
    """
    # ── Strategy 1: macOS textutil ────────────────────────────────────────
    try:
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = Path(tmp.name)

        result = subprocess.run(
            ["textutil", "-convert", "txt", str(path), "-output", str(tmp_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode == 0 and tmp_path.exists():
            text = tmp_path.read_text(encoding="utf-8", errors="replace").strip()
            tmp_path.unlink(missing_ok=True)
            if text:
                logger.info("Loaded DOC via textutil: %s (%d chars)", path.name, len(text))
                return RawDocument(
                    text=text,
                    source_file=path.name,
                    source_type=source_type,
                    file_format="doc",
                )
        tmp_path.unlink(missing_ok=True)
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass  # textutil not available — try next strategy

    # ── Strategy 2: python-docx (works if file is really .docx) ──────────
    try:
        return load_docx(path, source_type=source_type)
    except Exception:
        pass

    # ── Strategy 3: Helpful error ─────────────────────────────────────────
    raise ValueError(
        f"Cannot read binary .doc file: '{path.name}'.\n"
        "Convert it to .docx first:\n"
        "  • macOS: open in Word/LibreOffice → Save As .docx\n"
        "  • CLI:   libreoffice --headless --convert-to docx \"" + str(path) + "\"\n"
        "  • CLI:   soffice --headless --convert-to docx \"" + str(path) + "\""
    )


# ── TXT ───────────────────────────────────────────────────────────────────────

def load_txt(path: Path, source_type: str = "article") -> RawDocument:
    """Load a plain-text file."""
    text = path.read_text(encoding="utf-8", errors="replace")
    logger.info("Loaded TXT: %s (%d chars)", path.name, len(text))

    return RawDocument(
        text=text,
        source_file=path.name,
        source_type=source_type,
        file_format="txt",
    )


# ── Directory scanner ─────────────────────────────────────────────────────────

LOADER_MAP = {
    ".pdf":  load_pdf,
    ".docx": load_docx,   # Office Open XML — fully supported
    ".doc":  load_doc,    # Legacy binary format — uses textutil on macOS
    ".txt":  load_txt,
}

SOURCE_TYPE_FOLDER_MAP = {
    "textbooks": "textbook",
    "guidelines": "guideline",
    "articles": "article",
}


def load_documents(raw_dir: Path) -> list[RawDocument]:
    """
    Recursively scan *raw_dir* and load all supported documents.

    Sub-folder names are used to infer source_type:
        raw/textbooks/  → source_type = "textbook"
        raw/guidelines/ → source_type = "guideline"
        raw/articles/   → source_type = "article"
        raw/<other>/    → source_type = "other"

    Args:
        raw_dir: Path to the raw data directory (e.g. data/raw/).

    Returns:
        List of RawDocument instances, one per file.
    """
    raw_dir = Path(raw_dir)
    if not raw_dir.exists():
        logger.warning("raw_dir does not exist: %s", raw_dir)
        return []

    documents: list[RawDocument] = []
    skipped = 0

    for file_path in sorted(raw_dir.rglob("*")):
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()
        loader = LOADER_MAP.get(suffix)
        if loader is None:
            logger.debug("Skipping unsupported file: %s", file_path.name)
            skipped += 1
            continue

        # Infer source_type from immediate parent folder name
        parent_folder = file_path.parent.name.lower()
        source_type = SOURCE_TYPE_FOLDER_MAP.get(parent_folder, "other")

        try:
            doc = loader(file_path, source_type=source_type)
            documents.append(doc)
        except Exception as exc:
            logger.error("Failed to load %s: %s", file_path, exc)

    logger.info(
        "Loaded %d documents from %s (skipped %d unsupported files)",
        len(documents),
        raw_dir,
        skipped,
    )
    return documents
